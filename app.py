"""
Nobel Biocare Proposal Generator
Web app for generating custom sales proposals, promotions, and new start deals.
"""

import os
import json
from datetime import datetime
from io import BytesIO
from functools import wraps

from flask import Flask, render_template, request, redirect, url_for, session, send_file, flash
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'nobel-proposal-2026')
APP_PASSWORD = os.environ.get('APP_PASSWORD', 'nobel2026')

# ── Brand colors (Nobel Biocare 2026) ──
BRAND_PRIMARY = '#000000'
BRAND_ACCENT = '#FED880'
BRAND_GREY = '#F2F1F0'
BRAND_BLUE = '#000000'  # kept for backward compat
BRAND_BLUE_RGB = RGBColor(0x00, 0x00, 0x00)
BRAND_ACCENT_RGB = RGBColor(0xFE, 0xD8, 0x80)
BRAND_GOLD_RGB = RGBColor(0xFE, 0xD8, 0x80)

# ── Discount Groups ──
DISCOUNT_GROUPS = {
    'Implants': ['NobelActive TiUltra Implants', 'NobelActive Implants', 'Nobel Biocare N1 TiUltra Implants',
                 'NobelParallel CC TiUltra Implants', 'NobelReplace CC TiUltra Implants', 'NobelPearl Ceramic Implants',
                 'NobelZygoma Implants'],
    'Abutments': ['Esthetic Abutments', 'Multi-Unit Abutments', 'Locator Abutments', 'GoldAdapt Abutments'],
    'Kits': ['Surgical Kits'],
    'Grafting': ['Regenerative - Grafting'],
    'Membranes': ['Regenerative - Membranes'],
    'Sutures': ['Sutures'],
    '3Shape': ['Capital - 3Shape'],
    'SprintRay': ['Capital - SprintRay'],
    'DEXIS': ['Capital - DEXIS'],
    'X-Guide': ['Capital - X-Guide'],
    'iCAM': ['Capital - iCAM'],
    'DTX': ['Capital - DTX'],
}

BONE_MILL_IDS = ['87465', '87466']
ADDON_TOOL_IDS = ['87467', '87468', '87712', '87713', '87714', '87715', '87716', '35425',
                  '87303', '87304', '87308', '87309', '87338', '87302']

# Kit IDs for new starts
PURESET_KIT_IDS = ['87294', '87295', '87296', '87293', '87305', '87306', '87307', '87469', '301267']

def get_discount_group(category):
    for group, cats in DISCOUNT_GROUPS.items():
        if category in cats:
            return group
    return 'Other'

PRODUCTS = None
def get_products():
    global PRODUCTS
    if PRODUCTS is None:
        with open(os.path.join(os.path.dirname(__file__), 'products.json'), 'r') as f:
            PRODUCTS = json.load(f)
    return PRODUCTS

def find_product(item_id):
    """Find a product by ID across all categories."""
    for category, prods in get_products().items():
        for p in prods:
            if p['id'] == item_id:
                return p, category
    return None, None

# ── Auth ──
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('logged_in'):
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        if request.form.get('password') == APP_PASSWORD:
            session['logged_in'] = True
            return redirect(url_for('index'))
        flash('Invalid password', 'error')
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('logged_in', None)
    return redirect(url_for('login'))

# ── Helpers for implant types ──
def get_implant_types():
    """Return list of implant type summaries for promo selection."""
    products = get_products()
    types = []
    implant_cats = [
        ('NobelActive TiUltra Implants', 'NobelActive TiUltra'),
        ('NobelActive Implants', 'NobelActive'),
        ('Nobel Biocare N1 TiUltra Implants', 'N1 TiUltra'),
        ('NobelParallel CC TiUltra Implants', 'NobelParallel CC TiUltra'),
        ('NobelReplace CC TiUltra Implants', 'NobelReplace CC TiUltra'),
        ('NobelPearl Ceramic Implants', 'NobelPearl'),
    ]
    for cat, name in implant_cats:
        if cat in products and products[cat]:
            types.append({'category': cat, 'name': name, 'price': products[cat][0]['price']})
    return types

def get_implant_prices_dict():
    """Return {category: price} for JS."""
    return {t['category']: t['price'] for t in get_implant_types()}

def get_kits():
    """Return PureSet kits list."""
    products = get_products()
    kits = []
    for p in products.get('Surgical Kits', []):
        if p['id'] in PURESET_KIT_IDS:
            kits.append(p)
    return kits

def get_bone_mills():
    products = get_products()
    return [p for p in products.get('Surgical Kits', []) if p['id'] in BONE_MILL_IDS]

def get_addon_tools():
    products = get_products()
    return [p for p in products.get('Surgical Kits', []) if p['id'] in ADDON_TOOL_IDS]

# ══════════════════════════════════════════════════════════
# STANDARD PROPOSALS
# ══════════════════════════════════════════════════════════

@app.route('/')
@login_required
def index():
    products = get_products()
    categories = list(products.keys())
    discount_groups = list(DISCOUNT_GROUPS.keys()) + ['Other']
    return render_template('index.html', categories=categories, products=products,
                         discount_groups=discount_groups, get_discount_group=get_discount_group,
                         bone_mills=get_bone_mills(), addon_tools=get_addon_tools())

@app.route('/generate', methods=['POST'])
@login_required
def generate():
    products = get_products()
    account_name = request.form.get('account_name', 'Customer')
    output_format = request.form.get('output_format', 'docx')
    rep_name = request.form.get('rep_name', '')
    notes = request.form.get('notes', '')
    
    group_discounts = {}
    for group in list(DISCOUNT_GROUPS.keys()) + ['Other']:
        key = f'discount_{group.replace(" ", "_").replace("-", "_")}'
        group_discounts[group] = float(request.form.get(key, 0))
    
    items = []
    for key, value in request.form.items():
        if key.startswith('qty_') and value and int(value) > 0:
            item_id = key.replace('qty_', '')
            qty = int(value)
            prod, category = find_product(item_id)
            if prod:
                dg = get_discount_group(category)
                disc = group_discounts.get(dg, 0)
                items.append({
                    'id': item_id, 'description': prod['description'],
                    'list_price': prod['price'], 'discount_pct': disc,
                    'discounted_price': prod['price'] * (1 - disc / 100),
                    'quantity': qty, 'category': category
                })
    
    if not items:
        flash('Please select at least one product', 'error')
        return redirect(url_for('index'))
    
    list_total = sum(i['list_price'] * i['quantity'] for i in items)
    final_total = sum(i['discounted_price'] * i['quantity'] for i in items)
    
    data = {
        'title': 'Custom Sales Offer',
        'account_name': account_name, 'rep_name': rep_name, 'notes': notes,
        'date': datetime.now().strftime('%B %d, %Y'),
        'items': items,
        'list_total': list_total,
        'discount_amount': list_total - final_total,
        'final_total': final_total
    }
    
    if output_format == 'pdf':
        return generate_pdf(data)
    return generate_docx(data)


# ══════════════════════════════════════════════════════════
# PROMOTIONS
# ══════════════════════════════════════════════════════════

def get_promotions():
    """Return all promotion definitions."""
    return [
        # Q1 Special
        {
            'id': 'xguide-allin', 'name': 'X-Guide "All-in" Bundle', 'category': 'q1',
            'deal_id': '24521', 'expires': '3/27/2026', 'savings_pct': '41',
            'summary': 'X-Guide NXT $25K + FastMap $10K + IconiX $0 + 60 implants at discount',
            'description': 'X-Guide NXT for $25,000 + FastMap for $10,000 + IconiX software at no charge + 60 NB Implants at customer\'s standard discount (40% max).',
            'includes': [
                'X-Guide NXT unit, instrument kit, 12 X-Clips w/ surgery credits, X-Mark SW Upgrade, training, planning SW',
                'FastMap with 4 FastMap Scan Bodies',
                'IconiX AI software at no charge ($2,900 MSRP)',
                '60 Nobel Biocare implants at your standing discount'
            ],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 60, 'kit_free': False,
            'equipment_price': 35000, 'equipment_msrp': 60399, 'equipment_name': 'X-Guide NXT + FastMap + IconiX',
            'extra_items': []
        },
        # Implant + Kit Bundles
        {
            'id': 'implant-kit-na-freehand', 'name': '25 Implants + NobelActive Freehand Kit', 'category': 'implant_bundles',
            'deal_id': '24579', 'savings_pct': '61',
            'summary': '25 implants at discount + NobelActive freehand kit with all drills & bone mills at no charge',
            'description': 'Buy 25 implants at your customer\'s standard discount (40% max), receive NobelActive freehand kit (includes all drills and bone mills) at no charge.',
            'includes': ['25 Nobel Biocare implants at standing discount', 'NobelActive Freehand Kit with all drills and bone mills — FREE'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 25, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'NobelActive Freehand Kit (all drills + bone mills)', 'price': 0, 'msrp': 14423}]
        },
        {
            'id': 'implant-kit-np-freehand', 'name': '25 Implants + NobelParallel Freehand Kit', 'category': 'implant_bundles',
            'deal_id': '24582', 'savings_pct': '58',
            'summary': '25 implants at discount + NobelParallel freehand kit at no charge',
            'description': 'Buy 25 implants at standard discount (40% max), receive NobelParallel freehand kit with all drills and bone mills at no charge.',
            'includes': ['25 Nobel Biocare implants at standing discount', 'NobelParallel Freehand Kit with all drills and bone mills — FREE'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 25, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'NobelParallel Freehand Kit (all drills + bone mills)', 'price': 0, 'msrp': 12734}]
        },
        {
            'id': 'implant-kit-nr-freehand', 'name': '25 Implants + NobelReplace Freehand Kit', 'category': 'implant_bundles',
            'deal_id': '24583', 'savings_pct': '64',
            'summary': '25 implants at discount + NobelReplace freehand kit at no charge',
            'description': 'Buy 25 implants at standard discount (40% max), receive NobelReplace freehand kit at no charge.',
            'includes': ['25 Nobel Biocare implants at standing discount', 'NobelReplace Freehand Kit with all drills — FREE'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 25, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'NobelReplace Freehand Kit (all drills)', 'price': 0, 'msrp': 16399}]
        },
        {
            'id': 'implant-kit-na-guided', 'name': '25 Implants + NobelActive Guided Kit', 'category': 'implant_bundles',
            'deal_id': '24425', 'savings_pct': '68',
            'summary': '25 implants at discount + NobelActive guided kit at no charge',
            'description': 'Buy 25 implants at standard discount (40% max), receive NobelActive guided kit with all drills and components at no charge.',
            'includes': ['25 Nobel Biocare implants at standing discount', 'NobelActive Guided Kit with all drills and components — FREE'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 25, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'NobelActive Guided Kit (all drills + components)', 'price': 0, 'msrp': 19920}]
        },
        {
            'id': 'implant-kit-np-guided', 'name': '25 Implants + NobelParallel Guided Kit', 'category': 'implant_bundles',
            'deal_id': '24426', 'savings_pct': '68',
            'summary': '25 implants at discount + NobelParallel guided kit at no charge',
            'description': 'Buy 25 implants at standard discount (40% max), receive NobelParallel guided kit at no charge.',
            'includes': ['25 Nobel Biocare implants at standing discount', 'NobelParallel Guided Kit with all drills and components — FREE'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 25, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'NobelParallel Guided Kit (all drills + components)', 'price': 0, 'msrp': 19160}]
        },
        {
            'id': 'implant-kit-nr-guided', 'name': '25 Implants + NobelReplace Guided Kit', 'category': 'implant_bundles',
            'deal_id': '24427', 'savings_pct': '71',
            'summary': '25 implants at discount + NobelReplace guided kit at no charge',
            'description': 'Buy 25 implants at standard discount (40% max), receive NobelReplace guided kit at no charge.',
            'includes': ['25 Nobel Biocare implants at standing discount', 'NobelReplace Guided Kit with all drills and components — FREE'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 25, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'NobelReplace Guided Kit (all drills + components)', 'price': 0, 'msrp': 22831}]
        },
        {
            'id': 'implant-kit-n1', 'name': '25 Implants + N1 Surgical Kit', 'category': 'implant_bundles',
            'deal_id': '24584', 'savings_pct': '58',
            'summary': '25 implants at discount + N1 kit with wrench, bone mills, OsseoShapers, OsseoDirectors',
            'description': 'Buy 25 implants at standard discount (40% max), receive N1 surgical kit at no charge.',
            'includes': ['25 Nobel Biocare implants at standing discount', 'N1 Surgical Kit with prosthetic wrench, bone mills, all OsseoShaper 2s, and 6 OsseoDirectors — FREE'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 25, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'N1 Surgical Kit (fully loaded)', 'price': 0, 'msrp': 12700}]
        },
        {
            'id': 'implant-motor-osseoset', 'name': '25 Implants + Osseoset 300 Motor', 'category': 'implant_bundles',
            'deal_id': '24430', 'savings_pct': '63',
            'summary': '25 implants at discount + Osseoset 300 (wired or wireless) at no charge',
            'description': 'Buy 25 implants at standard discount (40% max), receive Osseoset 300 at no charge.',
            'includes': ['25 Nobel Biocare implants at standing discount', 'Osseoset 300 motor (wired or wireless) — FREE'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 25, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'Osseoset 300 Motor', 'price': 0, 'msrp': 15878}]
        },
        {
            'id': 'implant-motor-kavo', 'name': '25 Implants + KaVo Motor', 'category': 'implant_bundles',
            'deal_id': '24431', 'savings_pct': '63',
            'summary': '25 implants at discount + KaVo ExpertSurg or MasterSurg at no charge',
            'description': 'Buy 25 implants at standard discount (40% max), receive KaVo motor at no charge.',
            'includes': ['25 Nobel Biocare implants at standing discount', 'KaVo ExpertSurg or MasterSurg motor — FREE'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 25, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'KaVo Motor (ExpertSurg/MasterSurg)', 'price': 0, 'msrp': 15822}]
        },
        # Specialty
        {
            'id': 'zygoma-startup', 'name': 'NobelZygoma TiUltra Startup', 'category': 'specialty',
            'deal_id': '24408', 'savings_pct': '36',
            'summary': '10 Zygoma implants + 10 MUAs at 25% off + full kit at 50% off',
            'description': '10 NobelZygoma TiUltra Implants at 25% off, 10 NobelZygoma MUA Xeal at 25% off, NobelZygoma TiUltra kit at 50% off.',
            'includes': ['10 NobelZygoma TiUltra Implants at 25% off', '10 NobelZygoma MUA Xeal at 25% off', 'NobelZygoma TiUltra Kit (108236) at 50% off'],
            'needs_implant_selection': False, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 0, 'kit_free': False,
            'equipment_price': 10827, 'equipment_msrp': 16914, 'equipment_name': 'Zygoma Startup Package',
            'extra_items': []
        },
        {
            'id': 'zygoma-upgrade', 'name': 'NobelZygoma TiUltra Upgrade', 'category': 'specialty',
            'deal_id': '24409', 'savings_pct': '31',
            'summary': '10 Zygoma implants + 10 MUAs at 25% off + upgrade tooling at 50% off',
            'description': 'For customers upgrading from TiUnite kit. 10 implants + 10 MUAs at 25% off + new tooling at 50% off.',
            'includes': ['10 NobelZygoma TiUltra Implants at 25% off', '10 NobelZygoma MUA Xeal at 25% off', 'Upgrade tooling (not in 88521) at 50% off'],
            'needs_implant_selection': False, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 0, 'kit_free': False,
            'equipment_price': 8737, 'equipment_msrp': 12734, 'equipment_name': 'Zygoma Upgrade Package',
            'extra_items': []
        },
        {
            'id': 'pearl-startup', 'name': 'NobelPearl Startup Bundle', 'category': 'specialty',
            'deal_id': '24586', 'savings_pct': '29',
            'summary': '15 NobelPearl implants at list + kit at no charge',
            'description': '15 NobelPearl implants at list price, NobelPearl kit at no charge.',
            'includes': ['15 NobelPearl Ceramic Implants at list price ($685/ea)', 'NobelPearl Tapered Surgery Kit (301267) — FREE ($8,500 MSRP)'],
            'needs_implant_selection': False, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 0, 'kit_free': False,
            'equipment_price': 9525, 'equipment_msrp': 13370, 'equipment_name': 'NobelPearl Bundle (15 implants + kit)',
            'extra_items': []
        },
        {
            'id': 'handpiece', 'name': 'Implant + Handpiece Bundle', 'category': 'specialty',
            'deal_id': '24585', 'savings_pct': '52',
            'summary': '15 implants at discount + surgical handpiece at no charge',
            'description': 'Buy 15 implants at standard discount (40% max), receive a surgical handpiece at no charge.',
            'includes': ['15 Nobel Biocare implants at standing discount', 'Surgical handpiece — FREE'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 15, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'Surgical Handpiece', 'price': 0, 'msrp': 5979}]
        },
        {
            'id': 'newgrad-freehand', 'name': 'New Grad — Freehand Bundle', 'category': 'specialty',
            'deal_id': '88067', 'savings_pct': '78',
            'summary': '10 implants at 50% off + 10 healing abutments + freehand kit + motor at 65% off + DTX Clinic',
            'description': 'New graduate package: 10 implants at 50% off, 10 healing abutments free, freehand kit free, Osseoset 300 at 65% off, 1yr DTX Studio Clinic free.',
            'includes': ['10 implants at 50% off', '10 Healing Abutments at no charge', 'Freehand Kit with all drills and bone mills — FREE', 'Osseoset 300 at 65% off', '1 year DTX Studio Clinic — FREE'],
            'needs_implant_selection': True, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 10, 'kit_free': False,
            'equipment_price': 6480, 'equipment_msrp': 29092, 'equipment_name': 'New Grad Freehand Package',
            'extra_items': []
        },
        {
            'id': 'newgrad-guided', 'name': 'New Grad — Guided Bundle', 'category': 'specialty',
            'deal_id': '88067', 'savings_pct': '78',
            'summary': '15 implants at 50% off + 15 healing abutments + guided kit + motor at 65% off + DTX Clinic',
            'description': 'New graduate package: 15 implants at 50% off, 15 healing abutments free, guided kit free, Osseoset 300 at 65% off, 1yr DTX Studio Clinic free.',
            'includes': ['15 implants at 50% off', '15 Healing Abutments at no charge', 'Guided Kit with all drills — FREE', 'Osseoset 300 at 65% off', '1 year DTX Studio Clinic — FREE'],
            'needs_implant_selection': True, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 15, 'kit_free': False,
            'equipment_price': 8017, 'equipment_msrp': 36482, 'equipment_name': 'New Grad Guided Package',
            'extra_items': []
        },
        # Tooling
        {
            'id': 'prosthetic-kit', 'name': 'Prosthetic Kit — 65% Off', 'category': 'tooling',
            'deal_id': '24587', 'savings_pct': '65',
            'summary': 'Prosthetic PureSet Basic at 65% off — $715',
            'description': 'Prosthetic PureSet Basic kit at 65% off list price.',
            'includes': ['Prosthetic PureSet Basic (87301) at 65% off — $715'],
            'needs_implant_selection': False, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 0, 'kit_free': False,
            'equipment_price': 715, 'equipment_msrp': 2043, 'equipment_name': 'Prosthetic Kit',
            'extra_items': []
        },
        {
            'id': 'torque-wrench', 'name': 'Prosthetic Torque Wrench — 65% Off', 'category': 'tooling',
            'deal_id': '9386', 'savings_pct': '65',
            'summary': 'Prosthetic Torque Wrench + Driver at 65% off — $299',
            'description': 'Prosthetic Torque Wrench + Driver at 65% off.',
            'includes': ['Prosthetic Torque Wrench + Driver at 65% off — $299'],
            'needs_implant_selection': False, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 0, 'kit_free': False,
            'equipment_price': 299, 'equipment_msrp': 854, 'equipment_name': 'Torque Wrench + Driver',
            'extra_items': []
        },
        {
            'id': 'tooling-refresh', 'name': 'Tooling Refresh — 40% Off', 'category': 'tooling',
            'deal_id': '19134', 'savings_pct': '40',
            'summary': '40% off list price on tooling refresh',
            'description': 'Tooling refresh at 40% off list price. Contact for specific items.',
            'includes': ['40% off list price on eligible tooling items'],
            'needs_implant_selection': False, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 0, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': 'Tooling Refresh',
            'extra_items': []
        },
        {
            'id': 'education-tuition', 'name': 'Education Tuition Bundle', 'category': 'tooling',
            'deal_id': '108889', 'savings_pct': '25',
            'summary': '$1,000 off per 5 implants at list or $3K of NB products when bundling course tuition',
            'description': 'Buy 5 implants at list or $3,000 of NB products to get $1,000 off course tuition. Stackable.',
            'includes': ['5 implants at list price OR $3,000 NB products', '$1,000 off course tuition', 'Stackable — buy 20 implants to get $4,000 off', 'Contact Natalie Navarrete in T&E for course registration'],
            'needs_implant_selection': True, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 5, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'Tuition Discount', 'price': -1000, 'msrp': 0}]
        },
        {
            'id': 'nobelvision-new', 'name': 'NobelVision — New User', 'category': 'tooling',
            'deal_id': '10290', 'savings_pct': '48',
            'summary': '20 implants at list + 2yrs NobelVision service — $11,070 total',
            'description': '20 implants at list price + 2 years NobelVision service.',
            'includes': ['20 implants at list price', '2 years NobelVision service — FREE'],
            'needs_implant_selection': True, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 20, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': '2yr NobelVision Service', 'price': 0, 'msrp': 10290}]
        },
        {
            'id': 'nobelvision-current', 'name': 'NobelVision — Current User', 'category': 'tooling',
            'deal_id': '10292', 'savings_pct': '55',
            'summary': '12 implants at list + 2yrs NobelVision service — $4,380 total',
            'description': '12 implants at list price + 2 years NobelVision service renewal.',
            'includes': ['12 implants at list price', '2 years NobelVision service renewal — FREE'],
            'needs_implant_selection': True, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 12, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': '2yr NobelVision Service', 'price': 0, 'msrp': 5280}]
        },
        # Digital
        {
            'id': 'dexis-imprevo', 'name': 'DEXIS Imprevo Bundle', 'category': 'digital',
            'deal_id': '24588', 'savings_pct': '47',
            'summary': '60 implants + Imprevo for $10K + laptop 50% off',
            'description': 'DEXIS Imprevo scanner for $10,000 + Dell laptop at 50% off ($1,943) + 60 implants at standard discount (40% max).',
            'includes': ['DEXIS Imprevo Scanner ($23,000 MSRP) for $10,000', 'Dell Precision 7680 laptop at 50% off — $1,943', '60 implants at standing discount'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 60, 'kit_free': False,
            'equipment_price': 10000, 'equipment_msrp': 23000, 'equipment_name': 'DEXIS Imprevo Scanner',
            'extra_items': [{'name': 'Dell Precision 7680 Laptop (50% off)', 'price': 1943, 'msrp': 3885}]
        },
        {
            'id': 'dexis-3800', 'name': 'DEXIS 3800 Bundle', 'category': 'digital',
            'deal_id': '24590', 'savings_pct': '49',
            'summary': '55 implants + 3800W for $6K + laptop 50% off',
            'description': 'DEXIS IS 3800W for $6,000 + laptop at 50% off + 55 implants at standard discount (40% max).',
            'includes': ['DEXIS IS 3800W Scanner ($18,000 MSRP) for $6,000', 'Dell Precision 7680 laptop at 50% off — $1,943', '55 implants at standing discount'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 55, 'kit_free': False,
            'equipment_price': 6000, 'equipment_msrp': 18000, 'equipment_name': 'DEXIS IS 3800W Scanner',
            'extra_items': [{'name': 'Dell Precision 7680 Laptop (50% off)', 'price': 1943, 'msrp': 3885}]
        },
        {
            'id': 'trios-6', 'name': 'TRIOS 6 Bundle', 'category': 'digital',
            'deal_id': '24596', 'savings_pct': '35',
            'summary': '70 implants + TRIOS 6 for $21,500 + laptop 50% off',
            'description': '3Shape TRIOS 6 for $21,500 + laptop at 50% off + 70 implants at standard discount (40% max).',
            'includes': ['TRIOS 6 ($28,800 MSRP) for $21,500', 'Dell Precision 15" Touchscreen at 50% off — $1,750', '70 implants at standing discount'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 70, 'kit_free': False,
            'equipment_price': 21500, 'equipment_msrp': 28800, 'equipment_name': 'TRIOS 6 Scanner',
            'extra_items': [{'name': 'Dell Precision Laptop (50% off)', 'price': 1750, 'msrp': 3500}]
        },
        {
            'id': 'trios-5', 'name': 'TRIOS 5 Bundle', 'category': 'digital',
            'deal_id': '24598', 'savings_pct': '37',
            'summary': '70 implants + TRIOS 5 for $18,500 + laptop 50% off',
            'description': '3Shape TRIOS 5 for $18,500 + laptop at 50% off + 70 implants at standard discount (40% max).',
            'includes': ['TRIOS 5 ($26,700 MSRP) for $18,500', 'Dell Precision 15" Touchscreen at 50% off — $1,750', '70 implants at standing discount'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 70, 'kit_free': False,
            'equipment_price': 18500, 'equipment_msrp': 26700, 'equipment_name': 'TRIOS 5 Scanner',
            'extra_items': [{'name': 'Dell Precision Laptop (50% off)', 'price': 1750, 'msrp': 3500}]
        },
        {
            'id': 'trios-3', 'name': 'TRIOS 3 Bundle', 'category': 'digital',
            'deal_id': '24600', 'savings_pct': '36',
            'summary': '60 implants + TRIOS 3 for $13,500 + laptop 50% off',
            'description': '3Shape TRIOS 3 for $13,500 + laptop at 50% off + 60 implants at standard discount (40% max).',
            'includes': ['TRIOS 3 ($18,000 MSRP) for $13,500', 'Dell Precision 15" Touchscreen at 50% off — $1,750', '60 implants at standing discount'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 60, 'kit_free': False,
            'equipment_price': 13500, 'equipment_msrp': 18000, 'equipment_name': 'TRIOS 3 Scanner',
            'extra_items': [{'name': 'Dell Precision Laptop (50% off)', 'price': 1750, 'msrp': 3500}]
        },
        {
            'id': 'trios-move-plus', 'name': 'TRIOS MOVE Plus+ Bundle', 'category': 'digital',
            'deal_id': '24602', 'savings_pct': '52',
            'summary': '30 implants + MOVE Plus+ for $3,500',
            'description': '3Shape MOVE Plus+ for $3,500 + 30 implants at standard discount (40% max).',
            'includes': ['TRIOS MOVE Plus+ ($7,200 MSRP) for $3,500', '30 implants at standing discount'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 30, 'kit_free': False,
            'equipment_price': 3500, 'equipment_msrp': 7200, 'equipment_name': 'TRIOS MOVE Plus+',
            'extra_items': []
        },
        {
            'id': 'xguide-nxt', 'name': 'X-Guide NXT Bundle', 'category': 'digital',
            'deal_id': '24610', 'savings_pct': '38',
            'summary': '75 implants + X-Guide NXT for $27,999 + IconiX free',
            'description': 'X-Guide NXT for $27,999 + IconiX software at no charge + 75 implants at standard discount (40% max).',
            'includes': ['X-Guide NXT ($39,999 MSRP) for $27,999', 'IconiX AI software ($2,900 MSRP) — FREE', '75 implants at standing discount'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 75, 'kit_free': False,
            'equipment_price': 27999, 'equipment_msrp': 39999, 'equipment_name': 'X-Guide NXT',
            'extra_items': [{'name': 'IconiX AI Software', 'price': 0, 'msrp': 2900}]
        },
        {
            'id': 'fastmap', 'name': 'FastMap Bundle', 'category': 'digital',
            'deal_id': '24612', 'savings_pct': '40',
            'summary': '55 implants + FastMap for $10,500',
            'description': 'FastMap for $10,500 + 55 implants at standard discount (40% max).',
            'includes': ['FastMap License ($17,500 MSRP) for $10,500', '4 FastMap Scan Bodies included', '55 implants at standing discount'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 55, 'kit_free': False,
            'equipment_price': 10500, 'equipment_msrp': 17500, 'equipment_name': 'FastMap',
            'extra_items': []
        },
        {
            'id': 'icam-core', 'name': 'ICam Core Bundle', 'category': 'digital',
            'deal_id': '24627', 'savings_pct': '32',
            'summary': '40 implants + ICam Core Bundle for $22,500',
            'description': 'ICam Core Bundle for $22,500 + 40 implants at standard discount (40% max).',
            'includes': ['ICam Core Bundle ($29,900 MSRP) for $22,500', '8 ICamBodies, 8 ICamRefs, 1yr warranty, virtual training & installation', '40 implants at standing discount'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 40, 'kit_free': False,
            'equipment_price': 22500, 'equipment_msrp': 29900, 'equipment_name': 'ICam Core Bundle',
            'extra_items': []
        },
        {
            'id': 'sprintray-pro2', 'name': 'SprintRay Pro 2 / MIDAS Bundle', 'category': 'digital',
            'deal_id': '24632', 'savings_pct': '38',
            'summary': '40 implants + Pro 2 or MIDAS for $7K + NanoCure for $2,400',
            'description': 'SprintRay Pro 2 or MIDAS for $7,000 + NanoCure for $2,400 + 40 implants at standard discount (40% max).',
            'includes': ['SprintRay Pro 2 or MIDAS ($10,995 MSRP) for $7,000', 'NanoCure ($3,599 MSRP) for $2,400', '40 implants at standing discount'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 40, 'kit_free': False,
            'equipment_price': 7000, 'equipment_msrp': 10995, 'equipment_name': 'SprintRay Pro 2 / MIDAS',
            'extra_items': [{'name': 'NanoCure', 'price': 2400, 'msrp': 3599}]
        },
        {
            'id': 'dtx-clinic', 'name': 'DTX Studio Clinic Bundle', 'category': 'digital',
            'deal_id': '24641', 'savings_pct': '40',
            'summary': '5 implants + DTX Studio Clinic at 40% off',
            'description': 'DTX Studio Clinic (single or team license) at 40% off + 5 implants at standard discount (40% max).',
            'includes': ['DTX Studio Clinic Single ($2,388 MSRP) for $1,433 OR Team ($4,887 MSRP) for $2,932', '5 implants at standing discount', 'AI Diagnostic Module also available at 40% off'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 5, 'kit_free': False,
            'equipment_price': 1433, 'equipment_msrp': 2388, 'equipment_name': 'DTX Studio Clinic (Single)',
            'extra_items': []
        },
        {
            'id': 'cbct-100', 'name': 'CBCT Rebate — 100 Implants', 'category': 'digital',
            'deal_id': '11697', 'savings_pct': '49',
            'summary': '100 implants at list + $30K CBCT rebate',
            'description': '100 implants at list price + $30,000 CBCT rebate.',
            'includes': ['100 implants at list price', '$30,000 CBCT rebate after purchase'],
            'needs_implant_selection': True, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 100, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'CBCT Rebate', 'price': -30000, 'msrp': 0}]
        },
        {
            'id': 'cbct-200', 'name': 'CBCT Rebate — 200 Implants', 'category': 'digital',
            'deal_id': '11694', 'savings_pct': '49',
            'summary': '200 implants at list + $60K CBCT rebate',
            'description': '200 implants at list price + $60,000 CBCT rebate.',
            'includes': ['200 implants at list price', '$60,000 CBCT rebate after purchase'],
            'needs_implant_selection': True, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 200, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'CBCT Rebate', 'price': -60000, 'msrp': 0}]
        },
        {
            'id': 'cbct-300', 'name': 'CBCT Rebate — 300 Implants', 'category': 'digital',
            'deal_id': '81402', 'savings_pct': '49',
            'summary': '300 implants at list + $90K CBCT rebate',
            'description': '300 implants at list price + $90,000 CBCT rebate.',
            'includes': ['300 implants at list price', '$90,000 CBCT rebate after purchase'],
            'needs_implant_selection': True, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 300, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'CBCT Rebate', 'price': -90000, 'msrp': 0}]
        },
        # Digital Add-ons
        {
            'id': 'dexis-training', 'name': 'DEXIS Onsite Training + 10 Implants', 'category': 'digital',
            'deal_id': '24592', 'savings_pct': '52',
            'summary': '10 implants at discount + DEXIS onsite training at no charge ($1,500 value)',
            'description': 'DEXIS IS Onsite Training at no charge ($1,500 MSRP) + 10 implants at standard discount (40% max).',
            'includes': ['DEXIS IS Onsite Training ($1,500 MSRP) — FREE', '10 implants at standing discount'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 10, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'DEXIS IS Onsite Training', 'price': 0, 'msrp': 1500}]
        },
        {
            'id': 'icambody-set', 'name': 'ICamBody Set + 15 Implants', 'category': 'digital',
            'deal_id': '24629', 'savings_pct': '56',
            'summary': '15 implants at discount + ICamBody set at no charge ($3,400 value)',
            'description': 'ICamBody Set at no charge ($3,400 MSRP) + 15 implants at standard discount (40% max).',
            'includes': ['ICamBody Set ($3,400 MSRP) — FREE', '15 implants at standing discount'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 15, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'ICamBody Set', 'price': 0, 'msrp': 3400}]
        },
        {
            'id': 'sprintray-training-pro2', 'name': 'Pro 2 Training + 10 Implants', 'category': 'digital',
            'deal_id': '24637', 'savings_pct': '57',
            'summary': '10 implants at discount + Pro 2 in-office training at no charge ($2,500 value)',
            'description': 'SprintRay Pro 2 In-office Training at no charge ($2,500 MSRP) + 10 implants at standard discount (40% max).',
            'includes': ['Pro 2 In-office Training ($2,500 MSRP) — FREE', '10 implants at standing discount'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 10, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'Pro 2 In-office Training', 'price': 0, 'msrp': 2500}]
        },
        {
            'id': 'sprintray-training-midas', 'name': 'MIDAS Virtual Training + 5 Implants', 'category': 'digital',
            'deal_id': '24638', 'savings_pct': '48',
            'summary': '5 implants at discount + MIDAS virtual training at no charge ($615 value)',
            'description': 'MIDAS Virtual Training at no charge ($615 MSRP) + 5 implants at standard discount (40% max).',
            'includes': ['MIDAS Virtual Training ($615 MSRP) — FREE', '5 implants at standing discount'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 5, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'MIDAS Virtual Training', 'price': 0, 'msrp': 615}]
        },
        {
            'id': 'sprintray-protection', 'name': 'Pro 2 / MIDAS Protection Plan + 10 Implants', 'category': 'digital',
            'deal_id': '24636', 'savings_pct': '55',
            'summary': '10 implants at discount + protection plan at no charge (up to $2,022 value)',
            'description': 'Pro 2 Protection Plan ($2,022 MSRP) or MIDAS Protection Plan ($1,628 MSRP) at no charge + 10 implants at standard discount (40% max).',
            'includes': ['Pro 2 Protection Plan ($2,022) OR MIDAS Protection Plan ($1,628) — FREE', '10 implants at standing discount'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 10, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': [{'name': 'Protection Plan (Pro 2 or MIDAS)', 'price': 0, 'msrp': 2022}]
        },
        {
            'id': 'sprintray-accessories', 'name': 'SprintRay Accessories 50% Off + 10 Implants', 'category': 'digital',
            'deal_id': '24639', 'savings_pct': '43',
            'summary': '10 implants at discount + 50% off SprintRay accessories (Arch Kit, Duo Kit, etc.)',
            'description': '50% off SprintRay accessories + 10 implants at standard discount (40% max). Includes Arch Kit, Duo Kit, and more.',
            'includes': ['50% off SprintRay accessories (Arch Kit, Duo Kit, etc.)', '10 implants at standing discount', 'See Jamar\'s Cap Equip Pricelist for full options'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 10, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': []
        },
        {
            'id': 'sprintray-resin', 'name': 'SprintRay Resin 50% Off + 5 Implants', 'category': 'digital',
            'deal_id': '24640', 'savings_pct': '42',
            'summary': '5 implants at discount + 50% off resin bottle (OnX Tough 2, Surgical Guide, etc.)',
            'description': '50% off SprintRay resin + 5 implants at standard discount (40% max). Includes OnX Tough 2, Surgical Guide Resin, and more.',
            'includes': ['50% off SprintRay resin bottle (OnX Tough 2, Surgical Guide Resin, etc.)', '5 implants at standing discount'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 5, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': '',
            'extra_items': []
        },
        # Regen
        {
            'id': 'regen-newstart', 'name': 'New Start Regen Bundle', 'category': 'regen',
            'deal_id': '81405', 'savings_pct': '70',
            'summary': 'Bone + membrane + plugs for $397',
            'description': 'New customer regen starter: 2 creos xenoprotect or 1 RTMCollagen, 2 bone grafts, 1 Cytoplast plug box for $397 total.',
            'includes': ['2 creos xenoprotect 15x20mm OR 1 RTMCollagen 15x20mm', '2 allo.gain 1cc OR enCore 1cc OR xenogain 0.5g', '1 Cytoplast RTMPlug wound dressings (10/box)', 'Grand total: $397'],
            'needs_implant_selection': False, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 0, 'kit_free': False,
            'equipment_price': 397, 'equipment_msrp': 1300, 'equipment_name': 'New Start Regen Bundle',
            'extra_items': []
        },
        {
            'id': 'regen-sustainability', 'name': 'Regen Sustainability (Commitment)', 'category': 'regen',
            'deal_id': 'Signed form', 'savings_pct': '49',
            'summary': 'Purchase commitment with tiered discounts on creos, sutures, enCore, Cytoplast',
            'description': 'Yearly purchase commitment with maximum discounts: 49% creos, 25% sutures/NovaBone/Zcore/MicroDerm, 10% Pro-fix/Meisinger, 33% enCore, 40% Cytoplast wound dressings, 20% RPM/Membranes.',
            'includes': ['Up to 49% off creos products', 'Up to 25% off Sutures/NovaBone/Zcore/MicroDerm', 'Up to 33% off enCore', 'Up to 40% off Cytoplast wound dressings', 'Up to 20% off RPM/Membranes', 'Requires signed commitment form'],
            'needs_implant_selection': False, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 0, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': 'Regen Sustainability',
            'extra_items': []
        },
        {
            'id': 'osteogenics-6for5', 'name': 'Osteogenics 6 for 5', 'category': 'regen',
            'deal_id': '11426', 'savings_pct': '17',
            'summary': 'Buy 5 Osteogenics products, get 6 (16.67% savings)',
            'description': 'Osteogenics stackable offer: buy 5 get 6. 16.67% savings.',
            'includes': ['Buy 5 Osteogenics products, receive 6', '16.67% savings', 'Stackable with other offers'],
            'needs_implant_selection': False, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 0, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': 'Osteogenics 6 for 5',
            'extra_items': []
        },
        {
            'id': 'osteogenics-15for12', 'name': 'Osteogenics 15 for 12', 'category': 'regen',
            'deal_id': '11427', 'savings_pct': '20',
            'summary': 'Buy 12 Osteogenics products, get 15 (20% savings)',
            'description': 'Osteogenics stackable offer: buy 12 get 15. 20% savings.',
            'includes': ['Buy 12 Osteogenics products, receive 15', '20% savings', 'Stackable'],
            'needs_implant_selection': False, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 0, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': 'Osteogenics 15 for 12',
            'extra_items': []
        },
        {
            'id': 'osteogenics-40for30', 'name': 'Osteogenics 40 for 30', 'category': 'regen',
            'deal_id': '11428', 'savings_pct': '25',
            'summary': 'Buy 30 Osteogenics products, get 40 (25% savings)',
            'description': 'Osteogenics stackable offer: buy 30 get 40. 25% savings.',
            'includes': ['Buy 30 Osteogenics products, receive 40', '25% savings', 'Stackable'],
            'needs_implant_selection': False, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 0, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': 'Osteogenics 40 for 30',
            'extra_items': []
        },
        {
            'id': 'encore-3for2', 'name': 'enCore / Cytoplast MicroDerm / X-Gut — 3 for 2', 'category': 'regen',
            'deal_id': '12034', 'savings_pct': '33',
            'summary': 'Buy 2 get 3 on enCore, Cytoplast MicroDerm, or X-Gut (33% savings)',
            'description': 'Buy 2 get 3 on enCore, Cytoplast MicroDerm, or X-Gut products. 33% savings.',
            'includes': ['Buy 2, receive 3 (33% savings)', 'Applies to enCore, Cytoplast MicroDerm, X-Gut'],
            'needs_implant_selection': False, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 0, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': 'enCore/MicroDerm/X-Gut 3 for 2',
            'extra_items': []
        },
        {
            'id': 'cytoplast-wound', 'name': 'Cytoplast Wound Dressings — 40% Off', 'category': 'regen',
            'deal_id': '24292', 'savings_pct': '40',
            'summary': 'Cytoplast wound dressings at 40% off',
            'description': 'Cytoplast wound dressings at 40% off list price.',
            'includes': ['Cytoplast wound dressings at 40% off'],
            'needs_implant_selection': False, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 0, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': 'Cytoplast Wound Dressings',
            'extra_items': []
        },
        {
            'id': 'creos-10plus', 'name': 'creos 10+ Units — 30% Off', 'category': 'regen',
            'deal_id': '9175', 'savings_pct': '30',
            'summary': 'Buy 10+ units of creos at 30% off',
            'description': '10 or more units of creos products at 30% off.',
            'includes': ['10+ units of creos at 30% off'],
            'needs_implant_selection': False, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 0, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': 'creos 10+ Bulk',
            'extra_items': []
        },
        {
            'id': 'creos-25plus', 'name': 'creos 25+ Units — 40% Off', 'category': 'regen',
            'deal_id': '9174', 'savings_pct': '40',
            'summary': 'Buy 25+ units of creos at 40% off',
            'description': '25 or more units of creos products at 40% off.',
            'includes': ['25+ units of creos at 40% off'],
            'needs_implant_selection': False, 'needs_discount': False, 'needs_kit_selection': False,
            'implant_qty': 0, 'kit_free': False,
            'equipment_price': 0, 'equipment_msrp': 0, 'equipment_name': 'creos 25+ Bulk',
            'extra_items': []
        },
        {
            'id': 'regen-implant-combo', 'name': 'Regen + Implant Combo ($995)', 'category': 'regen',
            'deal_id': '108817', 'savings_pct': '40',
            'summary': '10 implants at discount + regen bundle for $995 (expires 4/30/2026)',
            'description': '10 implants at standing discount (max 40%) OR combo of implants/MIK/regen totaling $5K after discounts → regen bundle for $995. Expires 4/30/2026.',
            'includes': ['10 implants at standing discount (max 40%)', 'OR implants/MIK/regen combo totaling $5K after discounts', 'Regen bundle for $995', 'Max existing customer product discount of 40%'],
            'needs_implant_selection': True, 'needs_discount': True, 'needs_kit_selection': False,
            'implant_qty': 10, 'kit_free': False,
            'equipment_price': 995, 'equipment_msrp': 2500, 'equipment_name': 'Regen Bundle',
            'extra_items': [],
            'expires': '4/30/2026'
        },
    ]


@app.route('/promotions')
@login_required
def promotions():
    return render_template('promotions.html', promos=get_promotions())


@app.route('/promotions/<promo_id>')
@login_required
def promo_detail(promo_id):
    promo = next((p for p in get_promotions() if p['id'] == promo_id), None)
    if not promo:
        flash('Promotion not found', 'error')
        return redirect(url_for('promotions'))
    
    return render_template('promo_detail.html',
                         promo=promo,
                         promo_json=json.dumps(promo),
                         implant_types=get_implant_types(),
                         implant_prices_json=json.dumps(get_implant_prices_dict()),
                         available_kits=get_kits())


@app.route('/promotions/<promo_id>/generate', methods=['POST'])
@login_required
def generate_promo(promo_id):
    promo = next((p for p in get_promotions() if p['id'] == promo_id), None)
    if not promo:
        flash('Promotion not found', 'error')
        return redirect(url_for('promotions'))
    
    account_name = request.form.get('account_name', 'Customer')
    rep_name = request.form.get('rep_name', '')
    notes = request.form.get('notes', '')
    output_format = request.form.get('output_format', 'docx')
    discount_pct = min(float(request.form.get('discount_pct', 0)), 40)
    implant_type = request.form.get('implant_type', 'NobelActive TiUltra Implants')
    
    # Get implant price
    prices = get_implant_prices_dict()
    implant_price = prices.get(implant_type, 615)
    implant_qty = promo.get('implant_qty', 0)
    
    items = []
    
    # Implants
    if implant_qty > 0:
        impl_list = implant_price * implant_qty
        if promo.get('needs_discount'):
            impl_disc = impl_list * (1 - discount_pct / 100)
        else:
            impl_disc = impl_list  # e.g. CBCT at list
        items.append({
            'description': f'{implant_qty}x {implant_type.replace(" Implants", "")} Implants',
            'id': '-', 'list_price': implant_price, 'quantity': implant_qty,
            'discount_pct': discount_pct if promo.get('needs_discount') else 0,
            'discounted_price': impl_disc / implant_qty,
            'category': 'Implants'
        })
    
    # Equipment
    if promo.get('equipment_price', 0) > 0 or promo.get('equipment_msrp', 0) > 0:
        items.append({
            'description': promo.get('equipment_name', 'Equipment'),
            'id': promo.get('deal_id', '-'), 'list_price': promo.get('equipment_msrp', 0),
            'quantity': 1, 'discount_pct': 0,
            'discounted_price': promo.get('equipment_price', 0),
            'category': 'Equipment'
        })
    
    # Extra items
    for extra in promo.get('extra_items', []):
        items.append({
            'description': extra['name'], 'id': '-',
            'list_price': extra['msrp'], 'quantity': 1, 'discount_pct': 0,
            'discounted_price': extra['price'], 'category': 'Promo Item'
        })
    
    list_total = sum(i['list_price'] * i['quantity'] for i in items)
    final_total = sum(i['discounted_price'] * i['quantity'] for i in items)
    
    data = {
        'title': promo['name'],
        'account_name': account_name, 'rep_name': rep_name,
        'notes': notes or promo.get('description', ''),
        'date': datetime.now().strftime('%B %d, %Y'),
        'items': items,
        'list_total': list_total,
        'discount_amount': list_total - final_total,
        'final_total': final_total,
        'deal_id': promo.get('deal_id', '')
    }
    
    if output_format == 'pdf':
        return generate_pdf(data)
    return generate_docx(data)


# ══════════════════════════════════════════════════════════
# NEW SURGICAL STARTS
# ══════════════════════════════════════════════════════════

@app.route('/new-starts')
@login_required
def new_starts():
    products = get_products()
    implant_categories = {}
    for cat_name in ['NobelActive TiUltra Implants', 'NobelActive Implants', 'Nobel Biocare N1 TiUltra Implants',
                     'NobelParallel CC TiUltra Implants', 'NobelReplace CC TiUltra Implants', 'NobelPearl Ceramic Implants']:
        if cat_name in products:
            implant_categories[cat_name] = products[cat_name]
    
    return render_template('new_starts.html', implant_categories=implant_categories, kits=get_kits())


@app.route('/new-starts/generate', methods=['POST'])
@login_required
def generate_new_start():
    account_name = request.form.get('account_name', 'Customer')
    rep_name = request.form.get('rep_name', '')
    notes = request.form.get('notes', '')
    output_format = request.form.get('output_format', 'docx')
    discount_pct = min(float(request.form.get('discount_pct', 0)), 40)
    kit_id = request.form.get('kit_id', '')
    
    products = get_products()
    items = []
    total_implants = 0
    
    # Collect implant selections
    for key, value in request.form.items():
        if key.startswith('impl_') and value and int(value) > 0:
            item_id = key.replace('impl_', '')
            qty = int(value)
            prod, category = find_product(item_id)
            if prod:
                disc_price = prod['price'] * (1 - discount_pct / 100)
                items.append({
                    'description': prod['description'], 'id': item_id,
                    'list_price': prod['price'], 'quantity': qty,
                    'discount_pct': discount_pct, 'discounted_price': disc_price,
                    'category': category
                })
                total_implants += qty
    
    if total_implants < 10:
        flash('Minimum 10 implants required for New Surgical Start', 'error')
        return redirect(url_for('new_starts'))
    
    # Add kit at no charge
    kit_prod, kit_cat = find_product(kit_id)
    if kit_prod:
        items.append({
            'description': f'{kit_prod["description"]} (at no charge)',
            'id': kit_id, 'list_price': kit_prod['price'],
            'quantity': 1, 'discount_pct': 100, 'discounted_price': 0,
            'category': 'Kit (Free)'
        })
    
    list_total = sum(i['list_price'] * i['quantity'] for i in items)
    final_total = sum(i['discounted_price'] * i['quantity'] for i in items)
    
    data = {
        'title': 'New Surgical Start',
        'account_name': account_name, 'rep_name': rep_name,
        'notes': notes or f'New Surgical Start — {total_implants} implants at {discount_pct}% off + PureSet kit at no charge. Promo code: 81238',
        'date': datetime.now().strftime('%B %d, %Y'),
        'items': items,
        'list_total': list_total,
        'discount_amount': list_total - final_total,
        'final_total': final_total,
        'deal_id': '81238'
    }
    
    if output_format == 'pdf':
        return generate_pdf(data)
    return generate_docx(data)


# ══════════════════════════════════════════════════════════
# DOCUMENT GENERATION
# ══════════════════════════════════════════════════════════

def generate_docx(data):
    doc = Document()
    
    # Title
    title = doc.add_heading('Nobel Biocare', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = BRAND_BLUE_RGB
    
    subtitle = doc.add_heading(data.get('title', 'Custom Sales Offer'), level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph(f"Prepared for: {data['account_name']}")
    doc.add_paragraph(f"Date: {data['date']}")
    if data.get('rep_name'):
        doc.add_paragraph(f"Territory Manager: {data['rep_name']}")
    if data.get('deal_id'):
        doc.add_paragraph(f"Deal ID: {data['deal_id']}")
    doc.add_paragraph()
    
    # Table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Description', 'Qty', 'Item #', 'Discount', 'Unit Price']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    
    for item in data['items']:
        row = table.add_row().cells
        row[0].text = item['description']
        row[1].text = str(item['quantity'])
        row[2].text = str(item['id'])
        row[3].text = f"{item['discount_pct']:.0f}%"
        row[4].text = f"${item['discounted_price']:,.2f}"
    
    doc.add_paragraph()
    doc.add_paragraph(f"List Price Total: ${data['list_total']:,.2f}")
    doc.add_paragraph(f"Total Discount: -${data['discount_amount']:,.2f}")
    
    final_para = doc.add_paragraph()
    run = final_para.add_run(f"Your Price: ${data['final_total']:,.2f}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BRAND_BLUE_RGB
    
    savings_para = doc.add_paragraph()
    run = savings_para.add_run(f"You Save: ${data['discount_amount']:,.2f}")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x7a, 0x1a)
    
    if data.get('discount_amount') and data.get('list_total') and data['list_total'] > 0:
        pct = (data['discount_amount'] / data['list_total']) * 100
        doc.add_paragraph(f"Savings: {pct:.1f}%")
    
    if data.get('notes'):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Notes / Special Terms:").bold = True
        doc.add_paragraph(data['notes'])
    
    doc.add_paragraph()
    doc.add_paragraph("Thank you for choosing Nobel Biocare!")
    doc.add_paragraph("Prices valid for 30 days from date of proposal.")
    doc.add_paragraph("FOR INTERNAL USE ONLY. DO NOT DISTRIBUTE TO CUSTOMERS.")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"Nobel_{data.get('title','Proposal').replace(' ','_')}_{data['account_name'].replace(' ','_')}_{datetime.now().strftime('%Y%m%d')}.docx"
    return send_file(buffer, as_attachment=True, download_name=filename,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


def generate_pdf(data):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=24,
                                 textColor=colors.HexColor(BRAND_BLUE), alignment=1)
    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Heading2'], fontSize=16,
                                    textColor=colors.HexColor(BRAND_ACCENT), alignment=1)
    normal = styles['Normal']
    
    elements = []
    elements.append(Paragraph("Nobel Biocare", title_style))
    elements.append(Paragraph(data.get('title', 'Custom Sales Offer'), subtitle_style))
    elements.append(Spacer(1, 0.3*inch))
    
    elements.append(Paragraph(f"<b>Prepared for:</b> {data['account_name']}", normal))
    elements.append(Paragraph(f"<b>Date:</b> {data['date']}", normal))
    if data.get('rep_name'):
        elements.append(Paragraph(f"<b>Territory Manager:</b> {data['rep_name']}", normal))
    if data.get('deal_id'):
        elements.append(Paragraph(f"<b>Deal ID:</b> {data['deal_id']}", normal))
    elements.append(Spacer(1, 0.3*inch))
    
    # Table
    table_data = [['Description', 'Qty', 'Item #', 'Disc.', 'Unit Price']]
    for item in data['items']:
        table_data.append([
            item['description'], str(item['quantity']), str(item['id']),
            f"{item['discount_pct']:.0f}%", f"${item['discounted_price']:,.2f}"
        ])
    
    t = Table(table_data, colWidths=[3*inch, 0.5*inch, 0.9*inch, 0.5*inch, 0.9*inch])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(BRAND_BLUE)),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.lightgrey),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('TOPPADDING', (0, 1), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 5),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 0.3*inch))
    
    elements.append(Paragraph(f"List Price Total: ${data['list_total']:,.2f}", normal))
    elements.append(Paragraph(f"Total Discount: -${data['discount_amount']:,.2f}", normal))
    elements.append(Spacer(1, 0.1*inch))
    
    final_style = ParagraphStyle('Final', parent=normal, fontSize=14, fontName='Helvetica-Bold',
                                 textColor=colors.HexColor(BRAND_BLUE))
    elements.append(Paragraph(f"Your Price: ${data['final_total']:,.2f}", final_style))
    
    savings_style = ParagraphStyle('Savings', parent=normal, fontSize=12, fontName='Helvetica-Bold',
                                   textColor=colors.HexColor('#228B22'))
    elements.append(Paragraph(f"You Save: ${data['discount_amount']:,.2f}", savings_style))
    
    if data.get('notes'):
        elements.append(Spacer(1, 0.3*inch))
        bold_style = ParagraphStyle('BoldNote', parent=normal, fontSize=11, fontName='Helvetica-Bold')
        elements.append(Paragraph("Notes / Special Terms:", bold_style))
        elements.append(Paragraph(data['notes'], normal))
    
    elements.append(Spacer(1, 0.4*inch))
    footer = ParagraphStyle('Footer', parent=normal, fontSize=10, textColor=colors.grey)
    elements.append(Paragraph("Thank you for choosing Nobel Biocare!", footer))
    elements.append(Paragraph("Prices valid for 30 days from date of proposal.", footer))
    elements.append(Paragraph("FOR INTERNAL USE ONLY.", footer))
    
    doc.build(elements)
    buffer.seek(0)
    
    filename = f"Nobel_{data.get('title','Proposal').replace(' ','_')}_{data['account_name'].replace(' ','_')}_{datetime.now().strftime('%Y%m%d')}.pdf"
    return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/pdf')


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
