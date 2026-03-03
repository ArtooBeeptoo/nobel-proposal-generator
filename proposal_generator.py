"""
Nobel Biocare Proposal Generator
Generates PDF and Word proposals matching the official Nobel template
"""

import os
from io import BytesIO
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    Image, HRFlowable
)
from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Nobel Biocare brand colors
NOBEL_RED = colors.Color(0.89, 0.09, 0.22)  # #E31837
NOBEL_RED_RGB = RGBColor(227, 24, 55)  # For Word docs

# Asset paths
STATIC_DIR = os.path.dirname(os.path.abspath(__file__)) + '/static/images'
LOGO_PATH = STATIC_DIR + '/nobel-logo.png'
PRODUCT_COLLAGE_PATH = STATIC_DIR + '/image10.jpeg'

def get_category_display(category):
    """Map internal category names to display names"""
    mapping = {
        'implants': 'Implants',
        'healing_abutments': 'Prosthetics',
        'stock_abutments': 'Prosthetics',
        'temp_abutments': 'Prosthetics',
        'cover_screws': 'Prosthetics',
        'multiunit_abutments': 'Prosthetics',
        'locator_abutments': 'Prosthetics',
        'impression_taking': 'Prosthetics',
        'screws': 'Prosthetics',
        'replicas': 'Prosthetics',
        'procera_lab': 'Prosthetics',
        'prosthetic_drivers': 'Prosthetics',
        'regeneratives': 'Regenerative solutions',
        'sutures': 'Regenerative solutions',
        'surgical_kits': 'Tooling',
        'tooling': 'Tooling',
        'motors': 'Tooling',
        'dexis': 'Digital equipment',
        'trios': 'Digital equipment',
        'sprintray': 'Digital equipment',
        'icam': 'Digital equipment',
        'xguide': 'Digital equipment',
        'dtx_studio': 'Digital equipment',
    }
    return mapping.get(category, 'Other')

def calculate_totals(items):
    """Calculate list price, net price, and savings"""
    total_list = 0
    total_net = 0
    
    for item in items:
        qty = item.get('quantity', 1)
        price = item.get('price', 0)
        discount = item.get('discount', 0)
        total_list += price * qty
        total_net += price * (1 - discount/100) * qty
    
    return total_list, total_net, total_list - total_net

def group_items_by_display_category(items):
    """Group items by display category"""
    display_cats = {}
    
    for item in items:
        cat = item.get('category', 'Other')
        display = get_category_display(cat)
        if display not in display_cats:
            display_cats[display] = []
        display_cats[display].append(item)
    
    return display_cats


def generate_pdf(items, customer_name='', valid_through='', rep_name='', rep_title='', rep_phone='', rep_email='', notes=''):
    """Generate PDF proposal in Nobel Biocare format - fits on one page"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        leftMargin=0.5*inch, rightMargin=0.5*inch,
        topMargin=0.3*inch, bottomMargin=0.4*inch
    )
    
    styles = getSampleStyleSheet()
    
    # Determine sizing based on item count (shrink for many items)
    item_count = len(items)
    if item_count > 15:
        item_font_size = 7
        table_padding = 2
        spacer_size = 0.1
    elif item_count > 10:
        item_font_size = 8
        table_padding = 3
        spacer_size = 0.15
    else:
        item_font_size = 9
        table_padding = 4
        spacer_size = 0.2
    
    # Custom styles - LEFT ALIGNED
    headline_style = ParagraphStyle(
        'Headline', parent=styles['Normal'],
        fontSize=20, fontName='Helvetica-Bold',
        textColor=colors.black, spaceAfter=2, alignment=TA_LEFT
    )
    customer_style = ParagraphStyle(
        'Customer', parent=styles['Normal'],
        fontSize=16, fontName='Helvetica-Bold',
        textColor=colors.black, spaceAfter=6, alignment=TA_LEFT
    )
    red_text = ParagraphStyle(
        'RedText', parent=styles['Normal'],
        fontSize=10, fontName='Helvetica-Bold',
        textColor=NOBEL_RED, alignment=TA_LEFT
    )
    rep_name_style = ParagraphStyle(
        'RepName', parent=styles['Normal'],
        fontSize=9, fontName='Helvetica-Bold', alignment=TA_LEFT, leading=11
    )
    footer_style = ParagraphStyle(
        'Footer', parent=styles['Normal'],
        fontSize=6, textColor=colors.gray, alignment=TA_LEFT
    )
    
    story = []
    
    # No top bar - cleaner look
    
    # Header with logo (two-column layout)
    header_data = [[
        Paragraph("An exclusive offer for", headline_style),
        Image(LOGO_PATH, width=1.5*inch, height=0.5*inch) if os.path.exists(LOGO_PATH) else ''
    ]]
    header_table = Table(header_data, colWidths=[5*inch, 2.5*inch])
    header_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
    ]))
    story.append(header_table)
    
    # Customer name with red accent line
    story.append(Paragraph(customer_name or "_____________________", customer_style))
    story.append(HRFlowable(width=40, thickness=2, color=NOBEL_RED, spaceAfter=15))
    
    # Calculate totals
    total_list, total_net, savings = calculate_totals(items)
    display_cats = group_items_by_display_category(items)
    
    # Main content - two columns
    # Left: Pricing table, Right: Rep info + images
    
    # Build pricing table with item details
    table_data = []
    
    # Item style with dynamic font size
    item_style = ParagraphStyle(
        'ItemStyle', parent=styles['Normal'],
        fontSize=item_font_size, alignment=TA_LEFT
    )
    
    # Table header
    table_data.append([
        Paragraph('<b>Item</b>', item_style),
        Paragraph('<b>Qty</b>', item_style),
        Paragraph('<b>List</b>', item_style),
        Paragraph('<b>Disc</b>', item_style),
        Paragraph('<b>Net</b>', item_style),
    ])
    
    display_order = ['Implants', 'Prosthetics', 'Regenerative solutions', 
                     'Tooling', 'Digital equipment', 'Training and education']
    
    for display_cat in display_order:
        if display_cat in display_cats:
            cat_items = display_cats[display_cat]
            # Category header row
            cat_header = f'<font color="#E31837">►</font> <b>{display_cat}</b>'
            table_data.append([
                Paragraph(cat_header, item_style),
                '', '', '', ''
            ])
            
            # Individual items
            for item in cat_items:
                ref_num = item.get('id', '')
                desc = item.get('description', 'Item')[:35]  # Truncate long descriptions
                qty = item.get('quantity', 1)
                price = item.get('price', 0)
                discount = item.get('discount', 0)
                net_unit = price * (1 - discount / 100)
                net_extended = net_unit * qty
                
                ref_font_size = max(6, item_font_size - 1)
                item_text = f"<font size='{ref_font_size}'>#{ref_num}</font><br/>{desc}"
                table_data.append([
                    Paragraph(item_text, item_style),
                    str(qty),
                    f'${price:,.2f}',
                    f'{discount:.0f}%' if discount > 0 else '-',
                    f'${net_extended:,.2f}'
                ])
    
    if len(table_data) > 1:  # More than just header
        pricing_table = Table(table_data, colWidths=[2.8*inch, 0.4*inch, 0.7*inch, 0.4*inch, 0.7*inch], hAlign='LEFT')
        
        # Build style with category row highlighting - dynamic padding
        table_style = [
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), item_font_size),
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),  # First column left
            ('ALIGN', (1, 0), (-1, -1), 'RIGHT'),  # Rest right
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), table_padding),
            ('TOPPADDING', (0, 0), (-1, -1), table_padding),
            ('LINEBELOW', (0, 0), (-1, 0), 1, colors.black),  # Header line
            ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.95, 0.95, 0.95)),  # Header bg
        ]
        
        # Add background for category rows
        row_idx = 1
        for display_cat in display_order:
            if display_cat in display_cats:
                table_style.append(('BACKGROUND', (0, row_idx), (-1, row_idx), colors.Color(0.98, 0.96, 0.94)))
                table_style.append(('SPAN', (0, row_idx), (-1, row_idx)))
                row_idx += 1 + len(display_cats[display_cat])
        
        pricing_table.setStyle(TableStyle(table_style))
    else:
        pricing_table = Paragraph("<i>No items in quote</i>", styles['Normal'])
    
    # Totals section - compact
    totals_style = ParagraphStyle('TotalsStyle', parent=styles['Normal'], fontSize=10, fontName='Helvetica-Bold', alignment=TA_LEFT)
    totals_data = [
        [Paragraph('<b>Final sale price</b>', totals_style), f'${total_net:,.2f}'],
        [Paragraph('<b>Savings</b>', totals_style), f'${savings:,.2f}']
    ]
    totals_table = Table(totals_data, colWidths=[3.5*inch, 1*inch], hAlign='LEFT')
    totals_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
        ('BACKGROUND', (0, 1), (-1, 1), colors.Color(0.95, 0.93, 0.9)),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    
    # Add pricing table and totals
    story.append(pricing_table)
    story.append(Spacer(1, 4))
    story.append(totals_table)
    
    story.append(Spacer(1, 6))
    
    # Two-column layout: Left = validity/footer, Right = rep info + notes
    # Ultra-compact styling
    rep_small_style = ParagraphStyle('RepSmall', parent=styles['Normal'], fontSize=7, alignment=TA_LEFT, leading=9)
    validity_style = ParagraphStyle('Validity', parent=styles['Normal'], fontSize=7, alignment=TA_LEFT, leading=9)
    notes_style = ParagraphStyle('Notes', parent=styles['Normal'], fontSize=7, alignment=TA_LEFT, leading=9, textColor=colors.Color(0.3, 0.3, 0.3))
    
    # Build left column content (validity + footer) - ultra compact
    left_content = []
    left_content.append(Paragraph(f"<b>Offer valid through {valid_through or '___________'}.</b>", validity_style))
    left_content.append(Paragraph(
        "<b>This offer is valid only for the customer as listed and cannot be transferred, duplicated, or altered.</b>",
        validity_style
    ))
    left_content.append(Spacer(1, 3))
    left_content.append(Paragraph(
        "Nobel Biocare USA, LLC | 800-322-5001 | nobelbiocare.com",
        footer_style
    ))
    
    # Build right column content (rep info + notes)
    right_content = []
    if rep_name:
        right_content.append(Paragraph('<font color="#E31837"><b>This offer is presented by</b></font>', rep_small_style))
        right_content.append(Paragraph(f'<b>{rep_name}</b>', rep_name_style))
        if rep_title:
            right_content.append(Paragraph(rep_title, rep_small_style))
        if rep_phone:
            right_content.append(Paragraph(f'T {rep_phone}', rep_small_style))
        if rep_email:
            right_content.append(Paragraph(f'E {rep_email}', rep_small_style))
    
    if notes:
        if rep_name:
            right_content.append(Spacer(1, 4))
        right_content.append(Paragraph('<font color="#E31837"><b>Notes</b></font>', rep_small_style))
        # Handle multi-line notes
        for line in notes.split('\n'):
            if line.strip():
                right_content.append(Paragraph(line.strip(), notes_style))
    
    # Create two-column table
    left_cell = left_content if left_content else [Paragraph('', styles['Normal'])]
    right_cell = right_content if right_content else [Paragraph('', styles['Normal'])]
    
    # Wrap in nested tables for proper layout - zero padding
    left_table = Table([[item] for item in left_cell], colWidths=[4*inch])
    left_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))
    
    right_table = Table([[item] for item in right_cell], colWidths=[3*inch])
    right_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))
    
    two_col = Table([[left_table, right_table]], colWidths=[4*inch, 3.5*inch], hAlign='LEFT')
    two_col.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
        ('ALIGN', (1, 0), (1, 0), 'LEFT'),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))
    
    story.append(two_col)
    
    doc.build(story)
    buffer.seek(0)
    return buffer


def generate_docx(items, customer_name='', valid_through='', rep_name='', rep_title='', rep_phone='', rep_email=''):
    """Generate Word proposal in Nobel Biocare format"""
    doc = Document()
    
    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("An exclusive offer for")
    run.bold = True
    run.font.size = Pt(22)
    
    # Customer name
    customer_para = doc.add_paragraph()
    run = customer_para.add_run(customer_name or "_____________________")
    run.bold = True
    run.font.size = Pt(18)
    
    # Red accent line (using a table cell with red background)
    accent_table = doc.add_table(rows=1, cols=1)
    accent_table.rows[0].cells[0].width = Cm(2)
    accent_table.rows[0].height = Cm(0.1)
    for cell in accent_table.rows[0].cells:
        shading = cell._element.get_or_add_tcPr()
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'E31837')
        shading.append(shading_elm)
    
    doc.add_paragraph()  # Spacer
    
    # Calculate totals
    total_list, total_net, savings = calculate_totals(items)
    display_cats = group_items_by_display_category(items)
    
    # Pricing table with item details
    display_order = ['Implants', 'Prosthetics', 'Regenerative solutions', 
                     'Tooling', 'Digital equipment', 'Training and education']
    
    # Create table with 5 columns: Item, Qty, List, Disc, Net
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Header row
    header_row = table.rows[0]
    headers = ['Item', 'Qty', 'List', 'Disc', 'Net']
    for i, header in enumerate(headers):
        header_row.cells[i].paragraphs[0].add_run(header).bold = True
        if i > 0:
            header_row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    for display_cat in display_order:
        if display_cat in display_cats:
            cat_items = display_cats[display_cat]
            
            # Category header row
            cat_row = table.add_row()
            cell0 = cat_row.cells[0]
            run = cell0.paragraphs[0].add_run("► ")
            run.font.color.rgb = NOBEL_RED_RGB
            run.bold = True
            run2 = cell0.paragraphs[0].add_run(display_cat)
            run2.bold = True
            # Shade the category row
            for cell in cat_row.cells:
                shading = cell._element.get_or_add_tcPr()
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FAF5F0')
                shading.append(shading_elm)
            
            # Individual items
            for item in cat_items:
                item_row = table.add_row()
                ref_num = item.get('id', '')
                desc = item.get('description', 'Item')[:40]
                qty = item.get('quantity', 1)
                price = item.get('price', 0)
                discount = item.get('discount', 0)
                net_unit = price * (1 - discount / 100)
                net_extended = net_unit * qty
                
                # Item cell with ref# and description
                item_cell = item_row.cells[0]
                ref_run = item_cell.paragraphs[0].add_run(f"#{ref_num}\n")
                ref_run.font.size = Pt(8)
                ref_run.font.color.rgb = RGBColor(100, 100, 100)
                item_cell.paragraphs[0].add_run(desc)
                
                # Qty
                item_row.cells[1].text = str(qty)
                item_row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # List price
                item_row.cells[2].text = f"${price:,.2f}"
                item_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Discount
                item_row.cells[3].text = f"{discount:.0f}%" if discount > 0 else "-"
                item_row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Net extended
                item_row.cells[4].text = f"${net_extended:,.2f}"
                item_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()  # Spacer
    
    # Totals
    totals_table = doc.add_table(rows=2, cols=2)
    
    row1 = totals_table.rows[0]
    row1.cells[0].paragraphs[0].add_run("Final sale price").bold = True
    row1.cells[1].text = f"${total_net:,.2f}"
    row1.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    row2 = totals_table.rows[1]
    row2.cells[0].paragraphs[0].add_run("Savings").bold = True
    row2.cells[1].text = f"${savings:,.2f}"
    row2.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Highlight savings row
    for cell in row2.cells:
        shading = cell._element.get_or_add_tcPr()
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'F2EDE9')
        shading.append(shading_elm)
    
    doc.add_paragraph()  # Spacer
    
    # Rep info (only if provided)
    if rep_name:
        rep_para = doc.add_paragraph()
        run = rep_para.add_run("This offer is presented by")
        run.bold = True
        run.font.color.rgb = NOBEL_RED_RGB
        
        doc.add_paragraph()
        name_para = doc.add_paragraph()
        run = name_para.add_run(rep_name)
        run.bold = True
        run.font.size = Pt(14)
        
        if rep_title:
            doc.add_paragraph(rep_title)
        if rep_phone:
            doc.add_paragraph(f"T {rep_phone}")
        if rep_email:
            doc.add_paragraph(f"E {rep_email}")
        
        doc.add_paragraph()  # Spacer
    
    # Validity
    validity = doc.add_paragraph()
    run = validity.add_run(f"Offer valid through {valid_through or '___________'}.")
    run.bold = True
    
    disclaimer = doc.add_paragraph()
    run = disclaimer.add_run(
        "This offer is valid only for the customer as listed and cannot be transferred, duplicated, or altered."
    )
    run.bold = True
    
    doc.add_paragraph()  # Spacer
    
    # Footer
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Nobel Biocare USA, LLC. 22715 Savi Ranch Parkway, Yorba Linda, CA 92887; "
        "Toll free 800 322 5001; Technical support 888 725 7100\n"
        "MKT-5988.TMPLT Rev 00 © Nobel Biocare USA, LLC. All rights reserved. nobelbiocare.com"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
